from docx import Document
from docx.shared import Pt
doc = Document()

while True:
  exercise = input("What is the exercise (write 'end' to quit): ")
  formattedExercise = exercise.lower()
  containsSingle = "single" in formattedExercise
  print(containsSingle)
  if formattedExercise == "end":
    break
  while True:
    Weight = input("How much weight do you want to change? (Defaults to 5) Type 'bodyweight' if the exercise uses no weight: ")
    #elif Weight.lower() == "super-set":
      #break
    if Weight == "":
     Weight = float(5)
     break
    if Weight.lower() == "bodyweight":
      Weight = float(0)
      break
    elif not Weight.isdigit() or float(Weight) < 0:
      print("you can't do that")
      continue
    else:
      break
  paragraph = doc.add_paragraph(style='Normal')
  run21 = paragraph.add_run(exercise)
  run21.font.name = 'Courier New'
  run21.font.size = Pt(15)
  paragraph.alignment = 1
#def get_input_with_previous_values(prompt): 
  while True:
    #if Weight.lower() == "super-set":
      #Bench_Press_lbs = get_input_with_previous_values()
    if Weight == "bodyweight":
      Bench_Press_lbs = [0,0,0]
      break
    try:
      Bench_Press_lbs = list(map(float, input("Enter lbs for each set separated by space: ").split()))
      if any(lbs < 0 for lbs in Bench_Press_lbs):
          print("Please enter non-negative values.")
          continue
      break
    except ValueError:
      print("Please enter a valid number.")
  if len(Bench_Press_lbs) == 1:
    Bench_Press_lbs *= 3
  elif len(Bench_Press_lbs) == 2:
    Bench_Press_lbs.append(Bench_Press_lbs[-1])
  while True:
    if Weight == "bodyweight":
      how_many_reps = [7,7,7]
      break
    try:
      how_many_reps = list(
        map(int,
            input("Enter reps for each set separated by space: ").split()))
      if any(reps < 0 for reps in how_many_reps):
        print("Please enter non-negative values.")
        continue
      break
    except ValueError:
      print("Please enter a valid number.")
if len(how_many_reps) == 1:
    how_many_reps *= 3
elif len(how_many_reps) == 2:
    how_many_reps.append(how_many_reps[-1])
      # Check if the lengths of both lists match
if len(Bench_Press_lbs) != len(how_many_reps):
      print("Error: The number of pound inputs and rep inputs must be the same.")
else:
      # Process each set
    spam = 0 
    repsamount = 11
    for lbs, reps in zip(Bench_Press_lbs, how_many_reps, strict=True):
      if reps >= 12 - spam:
        print(containsSingle)
        if containsSingle:
          increase_lbs = lbs + float(Weight)
          paragraph1 = doc.add_paragraph(style='Normal')
          run1 = paragraph1.add_run(f"L {increase_lbs} lbs x {repsamount} reps R")
          run1.font.name = 'Courier New'
          run1.font.size = Pt(15)
          paragraph1.alignment = 1
        else: 
          increase_lbs = lbs + float(Weight)
          paragraph1 = doc.add_paragraph(style='Normal')
          run1 = paragraph1.add_run(f"{increase_lbs} lbs x {repsamount} reps")
          run1.font.name = 'Courier New'
          run1.font.size = Pt(15)
          paragraph1.alignment = 1
      elif 7 - spam <= reps <= 12 - spam:
        if containsSingle:
          keep_lbs = lbs
          paragraph2 = doc.add_paragraph(style='Normal')
          run2 = paragraph2.add_run(f"L {keep_lbs} lbs x {repsamount} reps R")
          run2.font.name = 'Courier New'
          run2.font.size = Pt(15)
          paragraph2.alignment = 1
        else:
          keep_lbs = lbs
          paragraph2 = doc.add_paragraph(style='Normal')
          run2 = paragraph2.add_run(f"{keep_lbs} lbs x {repsamount} reps")
          run2.font.name = 'Courier New'
          run2.font.size = Pt(15)
          paragraph2.alignment = 1
      elif 6 - spam >= reps >= 0:
        if containsSingle:
          lower_lbs = lbs - float(Weight)
          paragraph3 = doc.add_paragraph(style='Normal')
          run3 = paragraph3.add_run(f"L {lower_lbs} lbs x {repsamount} reps R")
          run3.font.name = 'Courier New'
          run3.font.size = Pt(15)
          paragraph3.alignment = 1
        else:
          lower_lbs = lbs - float(Weight)
          paragraph3 = doc.add_paragraph(style='Normal')
          run3 = paragraph3.add_run(f"{lower_lbs} lbs x {repsamount} reps")
          run3.font.name = 'Courier New'
          run3.font.size = Pt(15)
          paragraph3.alignment = 1

      spam += 2
      repsamount -= 2
doc.save('workout_results.docx')
